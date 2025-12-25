#!/usr/bin/env python3
"""
Script to convert DATA_FLOW_DIAGRAMS.md to DOCX format
Requires: pip install python-docx
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def markdown_to_docx(input_file, output_file):
    """Convert Markdown file to DOCX format"""
    
    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content by lines
    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            # Add empty paragraph
            doc.add_paragraph()
            continue
            
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('##### '):
            p = doc.add_heading(line[6:], level=5)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Handle horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Handle code blocks
        elif line.startswith('```'):
            # Skip code block markers, but handle content between them
            continue
            
        # Handle list items
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('2. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('3. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('4. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('5. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('6. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('7. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('8. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('9. '):
            p = doc.add_paragraph(line[3:], style='List Number')
        
        # Handle regular paragraphs
        else:
            # Handle bold text
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            # Handle italic text
            line = re.sub(r'\*(.*?)\*', r'\1', line)
            
            if line.strip():
                doc.add_paragraph(line)
    
    # Save the document
    doc.save(output_file)
    print(f"Successfully converted {input_file} to {output_file}")

if __name__ == "__main__":
    input_file = "DATA_FLOW_DIAGRAMS.md"
    output_file = "DATA_FLOW_DIAGRAMS.docx"
    
    try:
        markdown_to_docx(input_file, output_file)
    except Exception as e:
        print(f"Error converting file: {e}")
        print("Please install python-docx: pip install python-docx")
