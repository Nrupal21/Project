"""
Script to convert Project Report Markdown to Word Document
Requires: pip install python-docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def create_word_document():
    """Convert markdown project report to Word document"""
    
    # Read the markdown file
    markdown_file = "PROJECT_REPORT_DRAFT.md"
    if not os.path.exists(markdown_file):
        print(f"Error: {markdown_file} not found!")
        return
    
    with open(markdown_file, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Create Word document
    doc = Document()
    
    # Set up document styles
    setup_document_styles(doc)
    
    # Process content line by line
    lines = content.split('\n')
    
    for line in lines:
        process_line(doc, line)
    
    # Save the document
    output_file = "PROJECT_REPORT.docx"
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    
    return output_file

def setup_document_styles(doc):
    """Set up document styles for professional formatting"""
    
    # Configure normal style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

def process_line(doc, line):
    """Process each line of markdown and add to Word document"""
    
    line = line.strip()
    
    if not line:
        # Add empty paragraph
        doc.add_paragraph()
        return
    
    # Handle headings
    if line.startswith('# '):
        # Main heading (Chapter titles)
        p = doc.add_paragraph(line[2:])
        p.style = 'Heading 1'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=Pt(18))
        
    elif line.startswith('## '):
        # Sub-heading
        p = doc.add_paragraph(line[3:])
        p.style = 'Heading 2'
        set_paragraph_spacing(p, after=Pt(12))
        
    elif line.startswith('### '):
        # Sub-sub-heading
        p = doc.add_paragraph(line[4:])
        p.style = 'Heading 3'
        set_paragraph_spacing(p, after=Pt(6))
        
    elif line.startswith('#### '):
        # Fourth level heading
        p = doc.add_paragraph(line[5:])
        p.style = 'Heading 4'
        
    elif line.startswith('**') and line.endswith('**'):
        # Bold text
        p = doc.add_paragraph()
        run = p.add_run(line[2:-2])
        run.bold = True
        
    elif line.startswith('* '):
        # Bullet point
        p = doc.add_paragraph(line[2:], style='List Bullet')
        
    elif line.startswith('- '):
        # Bullet point (alternative)
        p = doc.add_paragraph(line[2:], style='List Bullet')
        
    elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
        # Numbered list
        p = doc.add_paragraph(line[3:], style='List Number')
        
    elif line.startswith('```'):
        # Code block - add as formatted text
        if line == '```' or line.startswith('```bash'):
            return  # Skip code block markers
        
    elif line.startswith('---'):
        # Horizontal line
        p = doc.add_paragraph()
        add_horizontal_line(p)
        
    else:
        # Regular paragraph
        if line and not line.startswith('!['):  # Skip image markdown
            doc.add_paragraph(line)

def set_paragraph_spacing(paragraph, before=None, after=None):
    """Set paragraph spacing"""
    paragraph_format = paragraph.paragraph_format
    if before:
        paragraph_format.space_before = before
    if after:
        paragraph_format.space_after = after

def add_horizontal_line(paragraph):
    """Add horizontal line to paragraph"""
    p = paragraph._p  # Get the underlying XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr, 'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 
                              'w:wordWrap', 'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 
                              'w:autoSpaceDN', 'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 
                              'w:spacing', 'w:ind', 'w:contextualSpacing', 'w:mirrorIndents', 
                              'w:suppressOverlap', 'w:jc', 'w:textDirection', 'w:textAlignment', 
                              'w:textboxTightWrap', 'w:outlineLvl', 'w:divId', 'w:cnfStyle', 
                              'w:rPr', 'w:sectPr', 'w:pPrChange')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_title_page(doc):
    """Add professional title page"""
    
    # Title
    title = doc.add_paragraph('ONLINE FOOD ORDERING SYSTEM')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run()
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Times New Roman'
    
    # Subtitle
    subtitle = doc.add_paragraph('A Project Report Submitted in Partial Fulfillment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run()
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'Times New Roman'
    
    subtitle2 = doc.add_paragraph('for the Award of Degree of')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2_run = subtitle2.add_run()
    subtitle2_run.font.size = Pt(14)
    subtitle2_run.font.name = 'Times New Roman'
    
    degree = doc.add_paragraph('MASTER OF COMPUTER APPLICATIONS')
    degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
    degree_run = degree.add_run()
    degree_run.font.size = Pt(14)
    degree_run.font.bold = True
    degree_run.font.name = 'Times New Roman'
    
    # Add spacing
    for _ in range(5):
        doc.add_paragraph()
    
    # Student info
    student = doc.add_paragraph('Submitted by:')
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    student_run = student.add_run()
    student_run.font.size = Pt(12)
    student_run.font.name = 'Times New Roman'
    
    name = doc.add_paragraph('[Your Name]')
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.add_run()
    name_run.font.size = Pt(12)
    name_run.font.bold = True
    name_run.font.name = 'Times New Roman'
    
    # Add more spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Department info
    dept = doc.add_paragraph('DEPARTMENT OF COMPUTER APPLICATIONS')
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_run = dept.add_run()
    dept_run.font.size = Pt(12)
    dept_run.font.bold = True
    dept_run.font.name = 'Times New Roman'
    
    university = doc.add_paragraph('[Your University Name]')
    university.alignment = WD_ALIGN_PARAGRAPH.CENTER
    university_run = university.add_run()
    university_run.font.size = Pt(12)
    university_run.font.name = 'Times New Roman'
    
    # Add more spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Date
    date = doc.add_paragraph('[Month Year]')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date.add_run()
    date_run.font.size = Pt(12)
    date_run.font.name = 'Times New Roman'
    
    # Page break
    doc.add_page_break()

def main():
    """Main function to create Word document"""
    print("📄 Creating Word Document from Project Report...")
    
    try:
        # Create document with title page
        doc = Document()
        setup_document_styles(doc)
        add_title_page(doc)
        
        # Read and process markdown content
        with open("PROJECT_REPORT_DRAFT.md", 'r', encoding='utf-8') as file:
            content = file.read()
        
        lines = content.split('\n')
        
        for line in lines:
            process_line(doc, line)
        
        # Save document
        output_file = "PROJECT_REPORT_DRAFT.docx"
        doc.save(output_file)
        
        print(f"✅ Word document created successfully: {output_file}")
        print(f"📁 File saved in: {os.path.abspath(output_file)}")
        print(f"📊 File size: {os.path.getsize(output_file)} bytes")
        
        return output_file
        
    except ImportError:
        print("❌ Error: python-docx package not installed!")
        print("📦 Install with: pip install python-docx")
        return None
    except FileNotFoundError:
        print("❌ Error: PROJECT_REPORT_DRAFT.md not found!")
        return None
    except Exception as e:
        print(f"❌ Error creating document: {str(e)}")
        return None

if __name__ == "__main__":
    main()
