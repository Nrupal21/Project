#!/usr/bin/env python3
"""
DOCX to Markdown Converter
Converts DOCX files to Markdown format while preserving formatting.
"""

import sys
import os
from docx import Document
from docx.shared import Inches
import re

def docx_to_markdown(docx_path, md_path=None):
    """
    Convert a DOCX file to Markdown format.
    
    Args:
        docx_path (str): Path to the input DOCX file
        md_path (str): Path to the output Markdown file (optional)
    
    Returns:
        str: The converted Markdown content
    """
    try:
        # Load the document
        doc = Document(docx_path)
        
        # Generate output filename if not provided
        if md_path is None:
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            md_path = f"{base_name}.md"
        
        markdown_content = []
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                markdown_content.append("")
                continue
            
            # Check heading level based on style
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.split()[-1]
                try:
                    level_num = int(level)
                    markdown_content.append(f"{'#' * level_num} {text}")
                except ValueError:
                    markdown_content.append(f"## {text}")
            else:
                # Handle bold and italic text
                for run in paragraph.runs:
                    if run.bold and run.italic:
                        text = text.replace(run.text, f"***{run.text}***")
                    elif run.bold:
                        text = text.replace(run.text, f"**{run.text}**")
                    elif run.italic:
                        text = text.replace(run.text, f"*{run.text}**")
                
                # Handle lists
                if text.startswith(('•', '-', '*')):
                    markdown_content.append(f"- {text[1:].strip()}")
                elif text[0].isdigit() and '.' in text[:5]:
                    markdown_content.append(f"{text}")
                else:
                    markdown_content.append(text)
        
        # Process tables
        for table in doc.tables:
            markdown_content.append("")
            
            # Table header
            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip())
            markdown_content.append("| " + " | ".join(header_row) + " |")
            markdown_content.append("|" + " | ".join(["-" * len(cell) for cell in header_row]) + " |")
            
            # Table data rows
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                markdown_content.append("| " + " | ".join(row_data) + " |")
            
            markdown_content.append("")
        
        # Join all content
        final_markdown = "\n".join(markdown_content)
        
        # Write to file
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(final_markdown)
        
        print(f"✅ Successfully converted {docx_path} to {md_path}")
        return final_markdown
        
    except Exception as e:
        print(f"❌ Error converting DOCX to Markdown: {str(e)}")
        return None

def main():
    """Main function to handle command line arguments."""
    if len(sys.argv) < 2:
        print("Usage: python docx_to_md.py <input_docx_file> [output_md_file]")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    md_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(docx_file):
        print(f"❌ Error: File '{docx_file}' not found.")
        sys.exit(1)
    
    result = docx_to_markdown(docx_file, md_file)
    
    if result:
        print(f"\n📄 Conversion completed successfully!")
        print(f"📊 Content length: {len(result)} characters")
    else:
        sys.exit(1)

if __name__ == "__main__":
    main()
